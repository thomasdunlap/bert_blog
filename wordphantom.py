import requests
from bs4 import BeautifulSoup
import re
import urllib.parse
from urllib.parse import urlparse
import argparse
import asyncio
import tensorflow as tf
from transformers import pipeline
import docx

def googleSearch(query):
    g_clean = [ ] #this is the list we store the search results
    url = 'https://www.google.com/search?client=ubuntu&channel=fs&q={}&ie=utf-8&oe=utf-8'.format(query)    
    try:
        html = requests.get(url)
        if html.status_code==200:
            soup = BeautifulSoup(html.text, 'lxml')
            a = soup.find_all('a') 
            for i in a:
                k = i.get('href')
                try:
                    m = re.search("(?P<url>https?://[^\s]+)", k)
                    n = m.group(0)
                    rul = n.split('&')[0]
                    domain = urlparse(rul)
                    if(re.search('google.com', domain.netloc)):
                        continue
                    else:
                        g_clean.append(rul)
                except:
                    continue
    except Exception as ex:
        print(str(ex))
    finally:
        return g_clean

def get_soup(url):
    html = requests.get(url)
    soup = BeautifulSoup(html.text, 'lxml')
    return soup

def get_text(query):
    urls = googleSearch(query)
    d = {}
    for url in urls[:5]:
        if url[-3:] == 'pdf':
            print(f"Gross. {url} is a pdf file.")
            continue
        else:
            soup = get_soup(url)
            d[url] = soup.get_text()

    return d

def write_docx(filepath, text, link, query):
    try:
        doc = docx.Document(filepath)
    except:
        doc = docx.Document()
    doc.add_heading(query, 1)
    doc.add_paragraph(text)
    doc.add_paragraph(link)
    doc.save(filepath)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(description="Scrape the googs!")
    parser.add_argument('query', type=str, help="Your google query")
    parser.add_argument('filepath', type=str, help="Word Document full filepath")
    args = parser.parse_args()
    summarize = pipeline("summarization")
    text = get_text(args.query)
    print(text, type(text))
    for url, t in text.items():
        
        try:
            remove_short = " ".join(line for line in t.split('\n') if len(line) > 20 or " [ " in line)
            print(f"Summarizing {url} ...")
            summary = summarize(remove_short, min_length=110, max_length=200)
            write_docx(args.filepath, summary[0]['summary_text'], url, args.query)
        except Exception as e:

            print(f"boo {url} ...")
            print(e)
    
            continue
    